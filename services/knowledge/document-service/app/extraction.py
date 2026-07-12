"""Lightweight resume extraction (no ML dependencies).

Uses pypdfium2 for PDF and python-docx for DOCX so the service runs within
small memory limits (e.g. a 512 MB free tier) with no model download or
cold-start load. Resumes are typed text documents, so ML-based layout/OCR
(previously Docling) is unnecessary here; the tradeoff is weaker handling of
scanned/image-only PDFs, which is acceptable for the MVP.

Privacy: never log file contents or extracted text.
"""

from pathlib import Path

import pypdfium2 as pdfium
from docx import Document


def _extract_pdf(path: Path) -> str:
    pdf = pdfium.PdfDocument(str(path))
    try:
        parts: list[str] = []
        for page in pdf:
            textpage = page.get_textpage()
            parts.append(textpage.get_text_range())
            textpage.close()
            page.close()
        return "\n\n".join(part.strip() for part in parts if part.strip())
    finally:
        pdf.close()


def _extract_docx(path: Path) -> str:
    document = Document(str(path))
    lines = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    # Resumes often lay out content in tables; include cell text too.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)
    return "\n".join(lines)


def extract_resume(path: Path) -> dict:
    """Extract plain text from a PDF or DOCX resume.

    Returns a dict with "text" and "markdown" (same value; the key shape is
    kept stable for the API response and callers).
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _extract_pdf(path).strip()
    elif suffix == ".docx":
        text = _extract_docx(path).strip()
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
    return {"text": text, "markdown": text}
